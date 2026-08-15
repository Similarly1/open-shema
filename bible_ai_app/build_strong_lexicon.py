import docx
import sys
import re
import json
import time

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

doc_p = r"C:\Users\adrie\kDrive\Documents\Livres Logos\DHG\Strong.docx"
t0 = time.time()
doc = docx.Document(doc_p)
t1 = time.time()
print(f"Loaded docx in {t1-t0:.2f}s, total paragraphs: {len(doc.paragraphs)}")

entries = {}
# Regex pattern for Strong's entries
# Examples: [[@HebrewStrongs:H2]]2. אַב : père
# or [[@GreekStrongs:G1]]1. Ἀλφα : Alpha
pattern = r'\[\[@(Hebrew|Greek)Strongs:([HG]\d+)\]\](\d+)\.\s*([^:]+?)\s*:\s*(.+)'

current_strong = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if not txt:
        continue
    
    m = re.search(pattern, txt)
    if m:
        lang_type = m.group(1) # Hebrew or Greek
        strong_code = m.group(2) # H2, G1, etc.
        strong_num = m.group(3)
        lemma = m.group(4).strip()
        definition = m.group(5).strip()
        
        # Standardize Strong code (e.g. H2 -> H0002 or keep H2)
        # Sg1910.csv uses H7225, H0430, G3972, G0652 (4 digits with leading zeros)
        prefix = strong_code[0] # 'H' or 'G'
        num_part = int(re.sub(r'\D', '', strong_code))
        std_code_4d = f"{prefix}{num_part:04d}"
        
        entries[std_code_4d] = {
            "code": std_code_4d,
            "raw_code": strong_code,
            "num": num_part,
            "lang": "hebrew" if prefix == "H" else "greek",
            "lemma": lemma,
            "definition": definition,
            "details": []
        }
        current_strong = std_code_4d
    elif current_strong:
        # Paragraphs following the entry provide details/glosses
        if not txt.startswith("[[@"):
            entries[current_strong]["details"].append(txt)
        else:
            current_strong = None

print(f"Extracted {len(entries)} Strong entries!")
hebrew_count = sum(1 for e in entries.values() if e['lang'] == 'hebrew')
greek_count = sum(1 for e in entries.values() if e['lang'] == 'greek')
print(f"  Hebrew entries: {hebrew_count}")
print(f"  Greek entries: {greek_count}")

# Sample test on H7225, H0430, G3972, G2424
for test_c in ["H7225", "H0430", "H1254", "G3972", "G2424", "G5547"]:
    if test_c in entries:
        print(f"\n[{test_c}] {entries[test_c]['lemma']} -> {entries[test_c]['definition']}")
        if entries[test_c]['details']:
            print(f"   Details: {' / '.join(entries[test_c]['details'][:2])}")
