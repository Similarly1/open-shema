import logging
logger = logging.getLogger(__name__)
import re
import docx
from core.reference_parser import normalize_reference

def extract_text_from_docx(file_path):
    """Extrait tout le texte brut d'un fichier .docx"""
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def clean_logos_fields(text):
    """Supprime les balises internes de formatage Logos comme {{field-on:Bible}} ou {{field-off:Bible}}"""
    text = re.sub(r'\{\{field-on:.*?\}\}', '', text)
    text = re.sub(r'\{\{field-off:.*?\}\}', '', text)
    # Nettoyage des balises de liens internes résiduelles
    text = re.sub(r'<Bible:\s*([^>]+)>', r'\1', text)
    return text.strip()

def chunk_by_logos_tags(text, doc_type, doc_name):
    """
    Découpe le texte en blocs basés sur les balises Logos.
    Gère [[@Bible:Ref]], [[@Headword:Mot]], [[@Topic:Sujet]], [[@Article:Titre]] et [[@Dictionary:Terme]].
    """
    # Pattern étendu pour capter n'importe quelle balise Logos
    pattern = r'\[\[@(Bible|Headword|Topic|Article|Dictionary):\s*(.*?)\s*\]\]'
    
    chunks = []
    matches = list(re.finditer(pattern, text))
    
    if not matches:
        return chunks
        
    for i, match in enumerate(matches):
        tag_type = match.group(1).strip() # 'Bible' ou 'Headword'
        reference = match.group(2).strip() # ex: 'Gen 1:1' ou 'suicide'
        
        start_idx = match.end()
        end_idx = matches[i+1].start() if i + 1 < len(matches) else len(text)
        
        chunk_text = text[start_idx:end_idx]
        
        # Ne pas nettoyer les balises {{field-on/off}} ici pour pouvoir filtrer le texte biblique à l'affichage
        # chunk_text = clean_logos_fields(chunk_text)
        
        if chunk_text:
            # Normalisation de la référence (ex: 'Jean 3:16' -> 'Joh 3:16') si c'est une Bible
            norm_ref = normalize_reference(reference) if tag_type == "Bible" else reference
            
            # Extraire livre, chapitre, verset pour requêtes avancées
            parts = norm_ref.split(" ")
            book = parts[0] if len(parts) > 0 else ""
            chapter_verse = parts[1] if len(parts) > 1 else ""
            
            if ":" in chapter_verse:
                chapter_parts = chapter_verse.split(":")
                chapter = chapter_parts[0]
                verse = chapter_parts[1]
            else:
                chapter = chapter_verse
                verse = ""
                
            # Calculer un verset propre pour le tri (si intervalle comme 3-4, on prend le premier chiffre)
            verse_val = ""
            if verse.isdigit():
                verse_val = int(verse)
            elif "-" in verse:
                first_part = verse.split("-")[0]
                if first_part.isdigit():
                    verse_val = int(first_part)
            else:
                verse_val = verse
                
            metadata = {
                "type": doc_type,      # "Bible", "Commentaire", ou "Dictionnaire"
                "name": doc_name,
                "reference": norm_ref,
                "book": book,
                "chapter": int(chapter) if chapter.isdigit() else chapter,
                "verse": verse_val,
                "tag_type": tag_type
            }
            chunks.append({
                "text": chunk_text,
                "metadata": metadata
            })
            
    return chunks

def process_document(file_path, doc_type, doc_name):
    text = extract_text_from_docx(file_path)
    chunks = chunk_by_logos_tags(text, doc_type, doc_name)
    return chunks

def has_logos_tags(file_path):
    """Vérifie rapidement si le document contient au moins une balise Logos"""
    try:
        doc = docx.Document(file_path)
        # Scan rapide des 100 premiers paragraphes
        for i, para in enumerate(doc.paragraphs):
            if i > 100:
                break
            if "[[@" in para.text:
                return True
        # Scan complet si non trouvé au début (reste très rapide en mémoire)
        for para in doc.paragraphs:
            if "[[@" in para.text:
                return True
    except Exception as e:
        logger.error("Erreur validation Logos :", e)
    return False
