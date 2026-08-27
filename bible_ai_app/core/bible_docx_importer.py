"""
Convertisseur et Importateur Universel de Bibles DOCX formatées Logos vers le format JSON de Bible AI.
Gère les documents Word balisés avec les ancres [[@Bible:Livre Chapitre:Verset]] et champs Logos.
"""

import os
import re
import json
import logging
from typing import Dict, Any, Optional, List, Tuple
from collections import defaultdict

import docx

from core.reference_parser import (
    BOOK_MAPPING, 
    REVERSE_BOOK_MAPPING, 
    strip_accents,
    get_standard_book_code,
    get_french_book_name,
    BOOKS_OT,
    BOOKS_NT,
    BOOKS_DEUTERO
)

logger = logging.getLogger(__name__)

STD_TO_USFM = {
    "Gen": "GEN", "Exo": "EXO", "Lev": "LEV", "Num": "NUM", "Deu": "DEU",
    "Jos": "JOS", "Jdg": "JDG", "Rut": "RUT", "1Sa": "1SA", "2Sa": "2SA",
    "1Ki": "1KI", "2Ki": "2KI", "1Ch": "1CH", "2Ch": "2CH", "Ezr": "EZR",
    "Neh": "NEH", "Est": "EST", "Job": "JOB", "Psa": "PSA", "Pro": "PRO",
    "Ecc": "ECC", "Sol": "SNG", "Isa": "ISA", "Jer": "JER", "Lam": "LAM",
    "Eze": "EZK", "Dan": "DAN", "Hos": "HOS", "Joe": "JOL", "Amo": "AMO",
    "Oba": "OBA", "Jon": "JON", "Mic": "MIC", "Nah": "NAM", "Hab": "HAB",
    "Zep": "ZEP", "Hag": "HAG", "Zec": "ZEC", "Mal": "MAL",
    "Mat": "MAT", "Mar": "MRK", "Luk": "LUK", "Joh": "JHN", "Act": "ACT",
    "Rom": "ROM", "1Co": "1CO", "2Co": "2CO", "Gal": "GAL", "Eph": "EPH",
    "Phi": "PHP", "Col": "COL", "1Th": "1TH", "2Th": "2TH", "1Ti": "1TI",
    "2Ti": "2TI", "Tit": "TIT", "Phm": "PHM", "Heb": "HEB", "Jam": "JAS",
    "1Pe": "1PE", "2Pe": "2PE", "1Jo": "1JN", "2Jo": "2JN", "3Jo": "3JN",
    "Jud": "JUD", "Rev": "REV",
    "Tob": "TOB", "Jdt": "JDT", "Wis": "WIS", "Sir": "SIR", "Bar": "BAR",
    "1Ma": "1MA", "2Ma": "2MA", "3Ma": "3MA", "4Ma": "4MA", "Man": "MAN"
}

def clean_verse_text(text: str) -> str:
    """Nettoie le texte d'un verset des balises Logos et résidus de formatage."""
    if not text:
        return ""
    # Retirer les balises de champs Logos {{field-on:...}} et {{field-off:...}}
    t = re.sub(r'\{\{field-on:.*?\}\}', '', text)
    t = re.sub(r'\{\{field-off:.*?\}\}', '', t)
    t = re.sub(r'\{\{[^}]+\}\}', '', t)
    # Retirer d'éventuelles balises de versets imbriquées [[@Bible:...]]
    t = re.sub(r'\[\[@[^\]]+\]\]', '', t)
    # Retirer le numéro de verset initial si présent
    t = re.sub(r'^[0-9]+\s*', '', t)
    # Remplacer les espaces insécables
    t = t.replace('\xa0', ' ')
    # Nettoyer les espaces multiples
    t = re.sub(r'\s+', ' ', t).strip()
    return t


class BibleDocxImporter:
    """Importateur et convertisseur haute performance pour Bibles Word DOCX (format Logos)."""

    @classmethod
    def is_logos_bible_docx(cls, docx_path: str) -> bool:
        """Vérifie rapidement si le document DOCX est une Bible balisée au format Logos."""
        if not os.path.exists(docx_path) or not docx_path.lower().endswith('.docx'):
            return False
        try:
            doc = docx.Document(docx_path)
            bible_tag_count = 0
            for i, para in enumerate(doc.paragraphs[:100]):
                if '[[@Bible:' in para.text:
                    bible_tag_count += 1
                    if bible_tag_count >= 2:
                        return True
            return bible_tag_count >= 1
        except Exception as e:
            logger.debug(f"[BibleDocxImporter] Erreur test Logos DOCX: {e}")
            return False

    @classmethod
    def inspect_bible_docx(cls, docx_path: str) -> Dict[str, Any]:
        """
        Inspecte un fichier DOCX Bible et extrait les métadonnées et chapitres pour l'assistant d'importation.
        """
        doc = docx.Document(docx_path)
        file_base = os.path.splitext(os.path.basename(docx_path))[0]

        title = ""
        author = ""
        year = ""

        # Détection du titre / auteur dans les premiers paragraphes
        for i, para in enumerate(doc.paragraphs[:12]):
            t = para.text.strip()
            if not t or '[[@Bible:' in t:
                continue
            # Ignorer si le paragraphe est simplement le nom d'un livre biblique (ex: Genèse, Exode...)
            if get_standard_book_code(t):
                continue
            if not title and (para.style.name.startswith('Heading') or any(kw in t.lower() for kw in ['bible', 'septante', 'testament', 'traduction', 'saintes ecritures'])):
                title = t
            elif not author and any(w in t.lower() for w in ['crampon', 'segond', 'darby', 'giguet', 'cahen', 'ostervald', 'chouraqui', 'auteur', 'traduction', 'abbé', 'chanoine', 'pasteur']):
                author = t
            elif not year and re.search(r'\b(1[5-9]\d{2}|20\d{2})\b', t):
                m_y = re.search(r'\b(1[5-9]\d{2}|20\d{2})\b', t)
                if m_y:
                    year = m_y.group(1)

        # Scan des livres et chapitres
        books_found = set()
        chapters_info = []

        all_canon = BOOKS_OT + BOOKS_NT + BOOKS_DEUTERO
        app_order = {code: i + 1 for i, (name, code, ch) in enumerate(all_canon)}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or '[[@Bible:' not in text:
                continue

            matches = re.findall(r'\[\[@Bible:([^:]+?)\s+(\d+):', text)
            for raw_b, chap_s in matches:
                b_code = get_standard_book_code(raw_b)
                if b_code:
                    books_found.add(b_code)

        # Création des chapitres d'inspection
        for b_code in sorted(books_found, key=lambda c: app_order.get(c, 999)):
            fr_name = get_french_book_name(b_code) or b_code
            scope = "AT" if any(b_code == c for _, c, _ in BOOKS_OT) else ("NT" if any(b_code == c for _, c, _ in BOOKS_NT) else "APO")
            chapters_info.append({
                "id": len(chapters_info) + 1,
                "title": fr_name,
                "book_code": b_code,
                "corpus_scope": scope,
                "source_type": "book_intro",
                "include": True,
                "size_chars": 1000
            })

        return {
            "title": title or file_base,
            "author": author or "",
            "year": year or "",
            "chapters": chapters_info,
            "total_books": len(books_found),
            "is_bible": True
        }

    @classmethod
    def parse_bible_docx(cls, docx_path: str) -> Dict[str, Dict[str, Dict[str, str]]]:
        """
        Extrait tous les livres, chapitres et versets d'un DOCX Logos.
        Retourne : { nom_livre_francais: { '1': { '1': 'Texte du verset...', '2': '...' } } }
        """
        doc = docx.Document(docx_path)
        bible_data: Dict[str, Dict[str, Dict[str, str]]] = defaultdict(lambda: defaultdict(dict))

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or '[[@Bible:' not in text:
                continue

            parts = re.split(r'\[\[@Bible:', text)
            for part in parts:
                if not part.strip():
                    continue

                m = re.match(
                    r'^([^:]+?)\s+(\d+):(\d+)\]\](?:\d+\s*)?(?:\{\{field-on:Bible\}\})?(.*?)(?:\{\{field-off:Bible\}\}|$)',
                    part.strip(),
                    re.DOTALL
                )
                if m:
                    book_raw, chap_s, verse_s, verse_text = m.groups()
                    chap_num = str(int(chap_s))
                    verse_num = str(int(verse_s))

                    clean_t = clean_verse_text(verse_text)
                    if not clean_t:
                        continue

                    std_code = get_standard_book_code(book_raw)
                    fr_name = get_french_book_name(std_code) or (book_raw.strip())

                    bible_data[fr_name][chap_num][verse_num] = clean_t

        # Convertir en dict standard et trier
        result: Dict[str, Dict[str, Dict[str, str]]] = {}
        for b_k, ch_dict in bible_data.items():
            if not ch_dict:
                continue
            result[b_k] = {}
            for ch_k in sorted(ch_dict.keys(), key=lambda x: int(x) if x.isdigit() else 0):
                if not ch_dict[ch_k]:
                    continue
                result[b_k][ch_k] = {}
                for v_k in sorted(ch_dict[ch_k].keys(), key=lambda x: int(x.split('-')[0]) if re.match(r'^\d+', x) else 0):
                    result[b_k][ch_k][v_k] = ch_dict[ch_k][v_k]

        return result

    @classmethod
    def import_bible_docx(
        cls, 
        docx_path: str, 
        custom_name: Optional[str] = None, 
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Importe et convertit une Bible DOCX complète dans data/bibles/
        Met à jour data/library.json et renvoie l'identifiant et les métadonnées.
        """
        from core.bible_json_loader import BibleJsonLoader

        all_books_list = BOOKS_OT + BOOKS_NT + BOOKS_DEUTERO
        app_order = {code: i + 1 for i, (name, code, ch) in enumerate(all_books_list)}

        parsed_bible = cls.parse_bible_docx(docx_path)
        if not parsed_bible:
            raise ValueError("Aucun verset biblique n'a pu être extrait du fichier DOCX Logos.")

        filename_base = os.path.splitext(os.path.basename(docx_path))[0]
        detected_name = custom_name or filename_base.replace("_", " ").replace("-", " ").title()
        version_code = (custom_metadata.get("version_code") if custom_metadata else None) or (custom_metadata.get("name") if custom_metadata else None) or detected_name[:4].upper().strip()
        version_fullname = (custom_metadata.get("title") if custom_metadata else None) or f"Bible {detected_name}"

        folder_clean = re.sub(r'[^\w\-_\. ]', '_', detected_name).strip().replace(" ", "_")
        bibles_dir = BibleJsonLoader.get_bibles_dir()
        dest_dir = os.path.join(bibles_dir, folder_clean)
        os.makedirs(dest_dir, exist_ok=True)

        # Nettoyage préalable des anciens fichiers JSON
        for old_f in os.listdir(dest_dir):
            if old_f.endswith('.json'):
                try:
                    os.remove(os.path.join(dest_dir, old_f))
                except Exception as err:
                    logger.warning(f"Impossible de supprimer {old_f}: {err}")

        saved_books_count = 0
        extra_idx = 67
        for raw_book_name, chapters_data in parsed_bible.items():
            std_code = get_standard_book_code(raw_book_name)
            fr_name = get_french_book_name(std_code) or raw_book_name
            if std_code in app_order:
                order_idx = app_order[std_code]
            else:
                order_idx = extra_idx
                extra_idx += 1
            usfm_code = STD_TO_USFM.get(std_code, std_code.upper() if std_code else "BOOK")

            book_obj = {
                "id": order_idx,
                "code": usfm_code,
                "name": fr_name,
                "version": version_code,
                "version_fullname": version_fullname,
                "total_chapters": len(chapters_data),
                "chapters": chapters_data
            }

            dest_filename = f"{order_idx:02d}_{usfm_code}_{fr_name}.json"
            dest_filepath = os.path.join(dest_dir, dest_filename)
            with open(dest_filepath, "w", encoding="utf-8") as fp:
                json.dump(book_obj, fp, ensure_ascii=False, indent=2)
            saved_books_count += 1

        lib_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "library.json")
        registry = {}
        if os.path.exists(lib_path):
            try:
                with open(lib_path, "r", encoding="utf-8") as f:
                    registry = json.load(f)
            except Exception:
                registry = {}

        meta = custom_metadata or {}
        bible_entry = {
            "title": meta.get("title") or detected_name,
            "author": meta.get("author") or "",
            "description": meta.get("description") or version_fullname,
            "year": meta.get("year") or "",
            "cover_path": meta.get("cover_path", None),
            "type": "Bible",
            "format": "json",
            "folder_name": folder_clean,
            "version_code": version_code,
            "total_books": saved_books_count,
            "embedding_model": "study_library",
            "active": True
        }

        registry[detected_name] = bible_entry
        os.makedirs(os.path.dirname(lib_path), exist_ok=True)
        with open(lib_path, "w", encoding="utf-8") as f:
            json.dump(registry, f, ensure_ascii=False, indent=2)

        BibleJsonLoader.clear_cache()
        logger.info(f"[BibleDocxImporter] Bible {detected_name} importée avec succès ({saved_books_count} livres).")

        return detected_name, bible_entry