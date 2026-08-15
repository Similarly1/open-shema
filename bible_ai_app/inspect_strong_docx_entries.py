import docx
import sys
import re

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

doc_p = r"C:\Users\adrie\kDrive\Documents\Livres Logos\DHG\Strong.docx"
doc = docx.Document(doc_p)

print(f"Total paragraphs in Strong.docx: {len(doc.paragraphs)}")

samples = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # Check for Strong number patterns like 0001, H0001, 1, 2, G0001, etc.
    if re.match(r'^[HG]?\d{1,5}\b', t) or re.match(r'^\d+\.?\s+', t):
        samples.append((i, t))
        if len(samples) >= 15:
            break

print("\n--- Detected Strong entry samples ---")
for idx, s in samples:
    print(f"[P{idx}] {s[:140]}")
